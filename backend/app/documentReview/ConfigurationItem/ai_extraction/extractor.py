"""
核心提取模块
实现配置项测试AI需求提取的主要功能
"""
import os
import time
import json
import logging
from docx import Document

from .api import query_ai_model
from .client import get_client
from .prompts import (
    EXTRACTION_PROMPT_TEMPLATE, 
    CONTINUATION_PROMPT_TEMPLATE,
    SPECIFIC_REQUIREMENT_PROMPT_TEMPLATE,
    SPECIFIC_CONTINUATION_PROMPT_TEMPLATE
)
from .chunking import chunk_document, chunk_document_with_overlap
from .utils import extract_json_from_text, merge_requirement_contents, extract_req_identifier

def ai_extract_requirements(file_path, model=None):
    """
    使用AI提取需求内容，支持滑动窗口和上下文回溯
    
    Args:
        file_path: Word文档路径
        model: 使用的AI模型名称
        
    Returns:
        提取的需求字典 {需求名称: 内容}
    """
    # 记录原始传入的模型名称以方便调试
    logging.debug(f"配置项AI提取需求原始传入模型: {model}")
    
    # 获取AI客户端
    client = get_client()
    if not client:
        logging.error("无法创建AI客户端，请检查配置")
        return {}
    
    # 记录开始时间
    start_time = time.time()
    
    # 从配置文件获取滑动窗口参数
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '../..', 'config.json')
    window_size = 3500  # 默认值
    overlap = 500  # 默认值
    enable_context = True  # 默认启用上下文回溯
    
    if os.path.exists(config_path):
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                extract_params = config.get("extraction_params", {})
                window_size = extract_params.get("window_size", 3500)
                overlap = extract_params.get("overlap", 500)
                enable_context = extract_params.get("enable_context", True)
                logging.info(f"使用配置的滑动窗口参数: 窗口大小={window_size}, 重叠={overlap}, 启用上下文={enable_context}")
        except Exception as e:
            logging.warning(f"读取滑动窗口配置失败: {str(e)}，使用默认配置")
    
    # 加载文档
    doc_load_start = time.time()
    doc = Document(file_path)
    logging.debug(f"文档加载耗时: {time.time() - doc_load_start:.2f}秒")
    
    # 提取文本内容
    text_start_time = time.time()
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    logging.debug(f"文档包含 {len(paragraphs)} 个段落，文本提取耗时: {time.time() - text_start_time:.2f}秒")
    
    # 分块处理文档，使用滑动窗口
    chunk_start_time = time.time()
    if enable_context:
        chunks, last_chapters = chunk_document_with_overlap(paragraphs, window_size, overlap)
        logging.debug(f"文档分成 {len(chunks)} 个滑动窗口，分块耗时: {time.time() - chunk_start_time:.2f}秒")
    else:
        chunks = chunk_document(paragraphs, max_tokens=window_size)  
        last_chapters = [None] * len(chunks)
        logging.debug(f"文档分成 {len(chunks)} 个固定块，分块耗时: {time.time() - chunk_start_time:.2f}秒")
    
    # 处理每个块并收集结果
    all_requirements = {}
    last_title = None
    
    for i, (chunk, last_chapter) in enumerate(zip(chunks, last_chapters)):
        logging.info(f"处理块 {i+1}/{len(chunks)}, 长度: {len(chunk)}")
        
        # 基于上下文构建提示
        if i == 0 or not enable_context or not last_title:
            # 第一个块或禁用上下文，使用原始提示
            prompt = EXTRACTION_PROMPT_TEMPLATE.format(chunk)
            logging.debug("使用初始提取提示")
        else:
            # 非第一个块，使用续取提示
            prompt = CONTINUATION_PROMPT_TEMPLATE.format(last_title, chunk)
            logging.info(f"使用续取提示，上一个提取的标题: {last_title}")
        
        try:
            # 查询AI模型
            response = query_ai_model(client, prompt, model)
            
            # 解析响应
            result = extract_json_from_text(response)
            if not result:
                logging.warning(f"块 {i+1} 解析失败，跳过")
                continue
                
            # 提取需求列表
            if result and "requirements" in result:
                requirements = result["requirements"]
                if isinstance(requirements, list):
                    # 处理提取到的需求
                    for req in requirements:
                        if isinstance(req, dict) and "title" in req:
                            title = req["title"]
                            chapter = req.get("chapter_number", "")
                            content = req.get("content", "")
                            identifier = req.get("identifier", "")
                            
                            # 如果内容非空，记录提取结果
                            if content and len(content.strip()) > 10:
                                # 记录提取到的需求
                                logging.info(f"提取到新需求: {title}, 章节: {chapter}, 标识: {identifier}")
                                
                                # 更新last_title用于下一次提取
                                if title:
                                    last_title = title
                                
                                # 使用标题作为key，内容作为value
                                if title not in all_requirements:
                                    all_requirements[title] = {
                                        "name": title,  # 使用name而不是title
                                        "chapter": chapter,  # 使用chapter而不是chapter_number
                                        "identifier": identifier,
                                        "content": content
                                    }
                                else:
                                    # 如果已存在，合并内容
                                    logging.debug(f"合并需求: {title}")
                                    existing = all_requirements[title]
                                    existing["content"] = merge_requirement_contents([existing["content"], content])
            else:
                logging.warning(f"块 {i+1} 未找到需求列表，跳过")
                continue
                
        except Exception as e:
            logging.error(f"处理块 {i+1} 时出错: {str(e)}")
    
    logging.info(f"AI需求提取完成，共提取到 {len(all_requirements)} 个需求，总耗时: {time.time() - start_time:.2f}秒")
    
    # 将字典转换为列表格式，便于前端处理
    requirements_list = list(all_requirements.values())
    
    return requirements_list

def ai_extract_specific_requirement(file_path, requirement_name, model=None):
    """
    使用AI提取特定需求的内容，支持滑动窗口处理
    
    Args:
        file_path: Word文档路径
        requirement_name: 需求名称
        model: 使用的AI模型名称
        
    Returns:
        需求内容字符串或None
    """
    # 记录原始传入的模型名称以方便调试
    logging.debug(f"配置项测试特定需求提取，原始传入模型: {model}")
    
    # 获取AI客户端
    client = get_client()
    if not client:
        logging.error("无法创建AI客户端，请检查配置")
        return None
    
    # 记录开始时间
    start_time = time.time()
    
    # 从配置文件获取滑动窗口参数
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '../..', 'config.json')
    window_size = 3500  # 默认值
    overlap = 500  # 默认值
    
    if os.path.exists(config_path):
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                extract_params = config.get("extraction_params", {})
                window_size = extract_params.get("window_size", 3500)
                overlap = extract_params.get("overlap", 500)
                logging.info(f"使用配置的滑动窗口参数: 窗口大小={window_size}, 重叠={overlap}")
        except Exception as e:
            logging.warning(f"读取滑动窗口配置失败: {str(e)}，使用默认配置")
    
    # 加载文档
    doc = Document(file_path)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    
    # 分块处理文档，使用滑动窗口
    chunks, _ = chunk_document_with_overlap(paragraphs, window_size, overlap)
    
    # 处理每个块并收集匹配结果
    matched_contents = []
    
    for i, chunk in enumerate(chunks):
        logging.info(f"特定需求搜索 - 处理块 {i+1}/{len(chunks)}")
        
        # 构建特定需求提取提示
        if i == 0 or not matched_contents:
            # 第一个块或尚未找到任何内容，使用初始提示
            prompt = SPECIFIC_REQUIREMENT_PROMPT_TEMPLATE.format(requirement_name, chunk)
        else:
            # 非第一个块，且已找到部分内容，使用续取提示
            prompt = SPECIFIC_CONTINUATION_PROMPT_TEMPLATE.format(requirement_name, chunk)
        
        try:
            # 查询AI模型
            response = query_ai_model(client, prompt, model)
            
            # 解析响应
            result = extract_json_from_text(response)
            
            if result and "requirement" in result:
                requirement = result["requirement"]
                if isinstance(requirement, dict) and "content" in requirement:
                    content = requirement["content"]
                    # 如果内容非空
                    if content and len(content.strip()) > 10:
                        logging.info(f"块 {i+1}: 找到相关内容，长度: {len(content)}")
                        matched_contents.append(content)
                        
                        # 顺便记录下标识符和章节，可能在其他场景有用
                        title = requirement.get("title", "")
                        chapter = requirement.get("chapter_number", "")
                        identifier = requirement.get("identifier", extract_req_identifier(content))
                        if identifier or chapter:
                            logging.info(f"提取到需求标识: {identifier}, 章节: {chapter}")
        except Exception as e:
            logging.error(f"处理块 {i+1} 失败: {str(e)}")
    
    # 合并所有匹配的内容
    final_content = None
    if matched_contents:
        final_content = merge_requirement_contents(matched_contents)
        logging.info(f"合并了 {len(matched_contents)} 个匹配片段，最终内容长度: {len(final_content)}")
    
    logging.debug(f"配置项特定需求搜索「{requirement_name}」完成，耗时: {time.time() - start_time:.2f}秒")
    
    return final_content

def ai_extract_named_requirements(file_path, requirement_names, model=None):
    """
    使用AI提取指定需求名称的内容
    
    Args:
        file_path: Word文档路径
        requirement_names: 需求名称列表
        model: 使用的AI模型名称
        
    Returns:
        提取的需求字典 {需求名称: 内容}
    """
    # 记录原始传入的模型名称以方便调试
    logging.debug(f"配置项命名需求提取，原始传入模型: {model}")
    
    result = {}
    start_time = time.time()
    
    for i, req_name in enumerate(requirement_names):
        logging.info(f"配置项提取需求 {i+1}/{len(requirement_names)}: {req_name}")
        content = ai_extract_specific_requirement(file_path, req_name, model)
        if content:
            result[req_name] = content
    
    logging.debug(f"配置项命名需求提取完成，耗时: {time.time() - start_time:.2f}秒，成功提取 {len(result)}/{len(requirement_names)} 个需求")
    
    return result

def ai_rematch_requirements(file_path, matched_requirements, model=None):
    """
    重新匹配未匹配到的需求
    
    Args:
        file_path: Word文档路径
        matched_requirements: 已匹配到的需求字典 {需求名称: 内容}
        model: 使用的AI模型名称
        
    Returns:
        合并后的需求字典 {需求名称: 内容}，包含原有的匹配需求和新匹配的需求
    """
    # 记录原始传入的模型名称以方便调试
    logging.debug(f"配置项AI重新匹配需求原始传入模型: {model}")
    
    # 导入重新匹配模块
    from .rematcher import rematch_requirements
    
    # 执行重新匹配
    new_requirements = rematch_requirements(file_path, matched_requirements, model=model)
    
    # 合并新旧需求
    all_requirements = {**matched_requirements, **new_requirements}
    
    logging.info(f"配置项需求重新匹配完成，原有{len(matched_requirements)}个需求，新增{len(new_requirements)}个需求，总计{len(all_requirements)}个")
    
    return all_requirements
