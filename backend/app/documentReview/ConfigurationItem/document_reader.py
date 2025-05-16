"""
需求文档读取模块
负责读取Word文档中的需求规格说明
"""
import os
import re
from docx import Document
import glob
import time

def check_word_document(folder_path):
    """
    判断文件夹下是否存在需求规格说明docx文件,如果是就返回对应的路径,如果不是就返回0
    
    参数:
        folder_path: 文件夹路径
        
    返回:
        匹配的文件路径或0
    """
    try:
        # 1、docx文档识别
        pattern = os.path.join(folder_path, "*需求规格说明*.docx")
        matching_files = glob.glob(pattern)
        if matching_files:
            return matching_files[0]
            
        # 2、doc文档识别
        pattern1 = os.path.join(folder_path, "*需求规格说明*.doc")
        matching_files1 = glob.glob(pattern1)
        if matching_files1:
            return matching_files1[0]
            
        # 如果没有匹配到任何文档
        return 0
    except Exception as e:
        print(f"发生异常: {e}")
        return 0





def read_xuqiu_wendang_document(file_path, requirement_names=None):
    """
    根据用户确认的需求名称列表提取需求
    参数：
        file_path: Word文档路径
        requirement_names: 用户确认的需求名称列表，如果为None或空列表，则自动提取
    返回：
        需求内容字典 {需求名称: 内容}
    """
    try:
        from docx import Document
        import re
        import time
        import os

        print(f"[DEBUG] read_xuqiu_wendang_document - 开始读取文件: {file_path}")
        print(f"[DEBUG] read_xuqiu_wendang_document - 需要提取的需求项数量: {len(requirement_names)}")
        
        # 设置全局环境变量，以便其他函数能够访问当前文档路径
        os.environ['CURRENT_DOC_PATH'] = file_path
        
        start_time = time.time()
        doc = Document(file_path)
        print(f"[DEBUG] 文档加载耗时: {time.time() - start_time:.2f}秒")
        
        # 提取所有段落文本
        text_extraction_start = time.time()
        text_list = [para.text.strip() for para in doc.paragraphs]
        len_temp = len(text_list)
        print(f"[DEBUG] 文档包含 {len_temp} 个段落，文本提取耗时: {time.time() - text_extraction_start:.2f}秒")
        
        requirement_dict = {}
        
        # 遍历需求名称
        total_items = len(requirement_names)
        for item_index, item in enumerate(requirement_names):
            item_start_time = time.time()
            print(f"[DEBUG] 正在处理第 {item_index+1}/{total_items} 个需求: {item}")
            
            if not isinstance(item, str):
                print(f"[WARNING] 跳过非字符串需求项: {item}")
                continue
                
            item_no_space = item.replace(' ', '')
            chapter_text = item  # 默认使用完整需求名称
            
            # 尝试提取章节号
            chapter_match = re.match(r'^([\d\.]+)\s+', item)
            chapter_number = chapter_match.group(1) if chapter_match else ""
            if chapter_number:
                # 提取章节号后的需求名称部分，用于更灵活的匹配
                chapter_text = item[len(chapter_number):].strip()
                print(f"[DEBUG] 需求「{item}」分离出章节号: {chapter_number} 和需求名称: {chapter_text}")
            else:
                print(f"[DEBUG] 需求「{item}」没有章节号")

            found = False
            
            # 设置搜索截止范围，防止无限搜索
            search_limit = min(len_temp, 1000)  # 限制搜索范围
            
            # 1. 先尝试完整匹配（包含章节号和需求名称）
            for index, text in enumerate(text_list[:search_limit]):
                # 添加搜索进度指示（每100个段落输出一次）
                if index % 100 == 0:
                    print(f"[DEBUG] 需求「{item}」搜索进度: {index}/{search_limit}段")
                
                # 查找需求项标题 - 完整匹配
                if item_no_space in text.replace(' ', ''):
                    found = True
                    print(f"[DEBUG] 在第{index}段找到需求项「{item}」(完整匹配)")
                    
                    # 提取需求内容
                    content = extract_requirement_content(text_list, index, len_temp, item)
                    if content:
                        requirement_dict[item] = content
                        break
            
            # 2. 如果完整匹配失败，尝试只用章节号匹配
            if not found and chapter_number:
                for index, text in enumerate(text_list[:search_limit]):
                    # 查找需求项标题 - 只用章节号匹配
                    if text.strip().startswith(chapter_number):
                        found = True
                        print(f"[DEBUG] 在第{index}段找到需求项「{item}」(使用章节号 {chapter_number} 匹配)")
                        
                        # 提取需求内容
                        content = extract_requirement_content(text_list, index, len_temp, item)
                        if content:
                            requirement_dict[item] = content
                            break
            
            # 3. 如果前两种方法都失败，尝试只用需求名称部分匹配
            if not found and chapter_text and len(chapter_text) > 5:
                for index, text in enumerate(text_list[:search_limit]):
                    # 查找需求项标题 - 只用需求名称部分匹配
                    if chapter_text.replace(' ', '') in text.replace(' ', ''):
                        found = True
                        print(f"[DEBUG] 在第{index}段找到需求项「{item}」(使用需求名称部分 {chapter_text} 匹配)")
                        
                        # 提取需求内容
                        content = extract_requirement_content(text_list, index, len_temp, item)
                        if content:
                            requirement_dict[item] = content
                            break
            
            # 如果所有方法都失败，使用占位符
            if not found:
                print(f"[WARNING] 未找到需求「{item}」，使用占位符内容")
                requirement_dict[item] = f"未找到需求「{item}」的详细内容"
            
            print(f"[DEBUG] 需求「{item}」处理完成，耗时: {time.time() - item_start_time:.2f}秒")
        
        total_time = time.time() - start_time
        print(f"[DEBUG] 共提取了{len(requirement_dict)}个需求项的内容，总耗时: {total_time:.2f}秒")
        for item, content in requirement_dict.items():
            print(f"[DEBUG] 需求「{item}」: 内容长度 {len(content)} 字符")
            
        return requirement_dict
    except Exception as e:
        print(f"读取需求文档时发生错误: {e}")
        import traceback
        traceback.print_exc()
        return {}

def extract_requirement_content(text_list, start_index, total_length, item_name):
    """
    提取需求内容的辅助函数，支持多种需求格式
    参数：
        text_list: 文档段落列表
        start_index: 匹配到需求名称的段落索引
        total_length: 文档总段落数
        item_name: 需求名称
    返回：
        提取的需求内容文本，如果未找到有效内容则返回None
    """
    from docx import Document
    import os
    
    print(f"[DEBUG] 开始从第{start_index}段提取需求「{item_name}」的内容")
    
    # 标准需求部分标识
    standard_parts = {
        'a)': '标识号',
        'b)': '说明',
        'c)': '进入条件',
        'd)': '输入',
        'e)': '输出',
        'f)': '处理',
        'g)': '性能',
        'h)': '约束'  # 匹配"约束"、"约束与限制"或"约束和限制"
    }
    
    # 关键词集合，用于识别需求内容
    key_terms = ['标识号', '说明', '进入条件', '输入', '输出', '处理', '性能', '约束']
    
    # 关键词变体映射，用于更灵活的匹配
    key_term_variants = {
        '标识号': ['标识号', '编号', 'id', '需求编号', 'req-'],
        '说明': ['说明', '描述', '功能描述', '概述', '简介'], 
        '进入条件': ['进入条件', '前置条件', '触发条件', '条件'],
        '输入': ['输入', '输入参数', '输入数据'],
        '输出': ['输出', '输出参数', '输出数据', '返回值'],
        '处理': ['处理', '处理过程', '处理逻辑', '实现过程', '步骤'],
        '性能': ['性能', '性能要求', '性能指标', '性能参数', '响应时间'],
        '约束': ['约束', '约束条件', '约束与限制', '约束和限制', '限制', '要求']
    }
    
    # 检查是否找到了有效的需求关键词
    found_valid_requirement = False
    
    # 记录遇到的分隔符位置，用于控制提取范围
    next_requirement_start = -1
    
    # 首先检查后续段落是否包含标准需求关键词
    for i in range(1, min(15, total_length - start_index - 1)):
        next_text = text_list[start_index + i].strip().lower()
        if any(key in next_text for key in key_terms):
            found_valid_requirement = True
            print(f"[DEBUG] 确认「{item_name}」是有效需求，找到关键词: {next_text[:20]}...")
            break
    
    # 预扫描，找出下一个需求的开始位置
    for i in range(start_index + 1, min(start_index + 200, total_length)):
        scan_text = text_list[i].strip()
        # 检测分隔符或新需求开始
        if (
            # 真正的分隔符
            scan_text.startswith('---') or
            
            # 新章节开始 - 确保不是处理步骤的编号
            (re.match(r'^[\d\.]+\s+[a-zA-Z\u4e00-\u9fa5]', scan_text) and 
             len(scan_text) < 80 and 
             # 确保不是当前需求
             not scan_text.startswith(item_name.split(' ')[0]) and
             # 特别要排除f)处理 部分中的数字编号步骤
             not any(scan_text.startswith(f"{n}.") for n in range(1, 10)) and
             # 性能和约束部分不应被识别为新需求
             not ('性能' in scan_text.lower() or '约束' in scan_text.lower()) and
             # 需要是形如"3.2.2 xxx功能"的标题，而不是处理步骤
             ('功能' in scan_text or '特性' in scan_text or '模块' in scan_text)
            )
        ):
            # 仍需确认这确实是新需求的开始
            # 进一步检查确认该段落不是当前需求的一部分
            is_valid_boundary = True
            
            # 检查前后文，如果前面有类似处理过程的特征，可能不是新需求
            if i > start_index + 5:  # 至少已经处理了几行
                prev_text = text_list[i-1].strip().lower()
                if (prev_text.startswith('f)') or 
                    '处理' in prev_text or 
                    prev_text.startswith('-') or
                    prev_text.endswith('：') or
                    prev_text.endswith(':')):
                    is_valid_boundary = False
            
            if is_valid_boundary:
                next_requirement_start = i
                print(f"[DEBUG] 预扫描发现下一个需求开始于第{i}段: {scan_text[:30]}...")
                break
    
    # 如果找到了下一个需求的位置，限制搜索范围，但至少要包含性能和约束部分的空间
    min_search_range = 30  # 确保至少有30段的搜索空间，以包含可能的性能和约束
    search_limit = max(min_search_range, 
                    (next_requirement_start - start_index - 1) if next_requirement_start > 0 
                    else min(100, total_length - start_index - 1))
    print(f"[DEBUG] 设置需求「{item_name}」的搜索范围为{search_limit}段")
    
    # 收集需求块内容
    block_content = [text_list[start_index]]  # 始终包含标题
    
    # 记录找到的需求部分
    requirement_parts_found = set()
    
    # 在搜索范围内查找所有与需求相关的内容
    i = 1
    while i <= search_limit:
        if start_index + i >= total_length:
            break
            
        next_text = text_list[start_index + i].strip()
        next_text_lower = next_text.lower()
        
        # 1. 检查是否遇到新章节（但确保不是太早停止）
        if i > 10 and re.match(r'^[\d\.]+\s+', next_text) and len(next_text) < 100:
            next_text_no_spaces = next_text.replace(' ', '')
            # 检查是否是一个章节标题（而不是需求内部的编号）
            # 避免将处理部分中的编号项误识别为新章节
            if (re.match(r'^[\d\.]+[a-zA-Z\u4e00-\u9fa5]', next_text_no_spaces) and
                not any(next_text.startswith(f"{n}.") for n in range(1, 10)) and  # 避免将1.2.3...格式的处理步骤误认为章节
                not ('性能' in next_text_lower or '约束' in next_text_lower or 
                     any(variant in next_text_lower for variant in key_term_variants['性能'] + key_term_variants['约束']))):
                print(f"[DEBUG] 在第{start_index+i}段遇到下一章节，停止提取: {next_text[:30]}...")
                break
        
        # 2. 识别标准需求格式的各部分
        is_requirement_part = False
        
        # 2.1 检查a)-h)格式
        for part_prefix, part_name in standard_parts.items():
            if next_text.startswith(part_prefix) or next_text.startswith(part_prefix.upper()):
                requirement_parts_found.add(part_name)
                is_requirement_part = True
                print(f"[DEBUG] 识别需求部分 {part_prefix}: {next_text[:30]}...")
                break
        
        # 2.2 检查直接使用关键词的格式
        if not is_requirement_part:
            for key, variants in key_term_variants.items():
                for variant in variants:
                    # 检查文本是否以变体开头或包含该变体
                    variant_lower = variant.lower()
                    if (next_text_lower.startswith(variant_lower) or 
                        (variant_lower in next_text_lower and len(next_text) < 150)):
                        
                        # 检查是否是标题式的格式
                        colon_index = max(next_text_lower.find('：'), next_text_lower.find(':'))
                        
                        if colon_index != -1 and colon_index < 30:  # 放宽冒号的位置限制
                            requirement_parts_found.add(key)
                            is_requirement_part = True
                            print(f"[DEBUG] 识别关键词「{variant}」格式的需求部分: {next_text[:30]}...")
                            break
                
                if is_requirement_part:
                    break
        
        # 3. 检查其他需求内容特征
        has_other_markers = (
            re.match(r'^[1-9]\d*\.', next_text) or  # 数字编号
            next_text.startswith('-') or  # 列表项
            not next_text or  # 空行
            (i < 8 and len(next_text) > 0)  # 标题后的前几行，放宽行数限制
        )
        
        # 4. 决定是否添加这一行
        if is_requirement_part or has_other_markers or found_valid_requirement:
            block_content.append(next_text)
            i += 1
            
            # 检查所有标准需求部分(a-h)是否都已找到
            part_prefixes = ['a)', 'b)', 'c)', 'd)', 'e)', 'f)', 'g)', 'h)']
            all_parts_found = True
            
            for part in part_prefixes:
                part_found = False
                for bc in block_content:
                    if bc.lower().startswith(part.lower()):
                        part_found = True
                        break
                
                if not part_found:
                    all_parts_found = False
                    break
            
            # 如果找到了所有标准部分(a-h)，尝试停止提取
            if all_parts_found and i > 20:  # 至少提取20段以确保完整捕获
                print(f"[DEBUG] 需求「{item_name}」已找到所有标准部分(a-h)，停止提取")
                break
        
        # 依然添加这一行，但记录警告
        else:
            if len(next_text) > 0:  # 忽略空行的警告
                print(f"[WARNING] 第{start_index+i}段可能不是需求内容，但仍添加: {next_text[:30]}...")
            block_content.append(next_text)
            i += 1
    
    # 找到的部分数量
    num_parts_found = len(requirement_parts_found)
    print(f"[DEBUG] 需求「{item_name}」找到了 {num_parts_found} 个标准部分: {', '.join(requirement_parts_found)}")
    
    # 处理表格内容
    try:
        # 获取当前文档的路径
        file_path = os.environ.get('CURRENT_DOC_PATH')
        if file_path and os.path.exists(file_path):
            doc = Document(file_path)
            item_no_space = item_name.replace(' ', '')
            
            for table_idx, table in enumerate(doc.tables):
                # 尝试检查表格是否与需求相关
                is_relevant_table = False
                
                # 1. 通过查找段落来检查
                for p_idx, paragraph in enumerate(doc.paragraphs):
                    p_text = paragraph.text.replace(' ', '')
                    if (item_no_space in p_text or 
                        any(key in paragraph.text.lower() for key in key_terms)):
                        if p_idx < len(doc.paragraphs) - 1:
                            next_p = doc.paragraphs[p_idx + 1]
                            if not next_p.text.strip() and table_idx < len(doc.tables):
                                is_relevant_table = True
                                break
                
                # 2. 通过表格内容检查
                if not is_relevant_table and table_idx < len(doc.tables):
                    try:
                        # 检查表格第一行是否包含需求相关词汇
                        first_row = table.rows[0]
                        first_row_text = " ".join([cell.text for cell in first_row.cells])
                        if (any(key in first_row_text.lower() for key in key_terms) or
                            item_no_space in first_row_text.replace(' ', '')):
                            is_relevant_table = True
                    except:
                        pass
                
                if is_relevant_table:
                    print(f"[DEBUG] 发现与需求「{item_name}」相关的表格")
                    table_content = ["需求相关表格:"]
                    
                    # 提取表格内容
                    for row_idx, row in enumerate(table.rows):
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            table_content.append(row_text)
                    
                    if len(table_content) > 1:  # 如果表格有内容
                        block_content.append("\n".join(table_content))
    except Exception as e:
        print(f"[WARNING] 处理表格时出错: {e}")
    
    # 单独检查是否提取了另一个需求的内容
    clean_content = []
    current_req_id = ""
    
    # 尝试从当前需求中提取标识号
    for text in block_content[:10]:  # 只检查前几行
        id_match = re.search(r'标识号[：:]\s*([\w\.-]+)', text, re.IGNORECASE)
        if id_match:
            current_req_id = id_match.group(1)
            print(f"[DEBUG] 需求「{item_name}」的需求标识号: {current_req_id}")
            break
    
    if current_req_id:
        # 处理内容，移除其他需求的部分
        in_current_req = True
        for text in block_content:
            # 检查是否遇到了新的需求标识号
            other_id_match = re.search(r'标识号[：:]\s*([\w\.-]+)', text, re.IGNORECASE)
            if other_id_match:
                other_req_id = other_id_match.group(1)
                if other_req_id != current_req_id:
                    print(f"[DEBUG] 发现其他需求标识号 {other_req_id}，停止接收内容")
                    in_current_req = False
            
            # 如果是当前需求的内容，或者没有检测到需求标识号，则添加内容
            if in_current_req:
                clean_content.append(text)
        
        if len(clean_content) < len(block_content):
            print(f"[DEBUG] 清理掉了 {len(block_content) - len(clean_content)} 段其他需求的内容")
            block_content = clean_content
    
    # 组合并处理需求块
    if block_content:
        content_text = "\n".join(block_content)
        content_length = len(content_text)
        
        # 内容评估 - 如果过短且无关键词，可能不是有效需求
        if content_length < 50 and num_parts_found == 0 and not found_valid_requirement:
            print(f"[WARNING] 需求「{item_name}」提取的内容过短({content_length}字符)且无关键词，可能不是有效需求")
            return None
            
        # 查找性能和约束部分的状态
        has_performance = '性能' in requirement_parts_found
        has_constraints = '约束' in requirement_parts_found
        
        if num_parts_found >= 4:  # 如果至少找到4个标准部分，认为是有效的需求
            # 输出性能和约束部分的提取状态
            if not has_performance:
                print(f"[WARNING] 需求「{item_name}」未找到性能部分")
            if not has_constraints:
                print(f"[WARNING] 需求「{item_name}」未找到约束部分")
                
            print(f"[DEBUG] 需求「{item_name}」内容提取成功，长度: {content_length} 字符，" +
                  f"找到标准部分: {num_parts_found} 个")
            return content_text
    
    return None

# ====== 从目录文件提取需求名称和章节号 ======
def extract_requirement_candidates(document_file_path, catalog_file_path):
    """
    从目录文件提取需求名称和章节号。
    
    参数:
        document_file_path: str, 主文档文件路径
        catalog_file_path: str, 目录文件路径
        
    目录文件格式应该类似图片中所示，包含章节号和需求名称（页码不是必需的）。
    如果章节下有子章节，则过滤掉该章节。
    返回：包含需求名称和章节号的对象列表 [{name, chapter}, ...]
    """
    from docx import Document
    import re
    import time

    # 设置环境变量，以便其他函数能够访问当前文档路径
    import os
    os.environ['CURRENT_DOC_PATH'] = document_file_path
    
    if not catalog_file_path:
        print("[ERROR] 必须提供目录文件！")
        return []
        
    print(f"[DEBUG] 使用目录文件提取需求点: {catalog_file_path}")
    
    total_start_time = time.time()
    
    # 从目录文件读取内容
    try:
        doc_load_start = time.time()
        doc = Document(catalog_file_path)
        print(f"[DEBUG] 目录文件加载耗时: {time.time() - doc_load_start:.2f}秒")
    except Exception as e:
        print(f"[ERROR] 打开目录文件时出错: {str(e)}")
        return []
    
    # 存储章节号和名称
    chapter_entries = []
    
    # 计时提取段落
    para_start_time = time.time()
    paragraphs = [para for para in doc.paragraphs if para.text.strip()]
    print(f"[DEBUG] 目录文件包含 {len(paragraphs)} 个非空段落，提取耗时: {time.time() - para_start_time:.2f}秒")
    
    # 计时处理段落
    process_start_time = time.time()
    
    for i, para in enumerate(paragraphs):
        if i % 50 == 0:  # 每处理50个段落输出一次进度
            print(f"[DEBUG] 正在处理段落 {i+1}/{len(paragraphs)}...")
            
        text = para.text.strip()
        if not text:
            continue
            
        # 匹配章节号和名称
        # 格式如: "3.2.1. 用户身份验证功能    2" 或 "3.2.1 用户身份验证功能    2" 或 "3.2.1 用户身份验证功能"
        # 章节号可以有多个层级 (如 3.2.1.4.5)，页码是可选的
        match = re.match(r'^(\d+(?:\.\d+)*(?:\.)?)\s+([^\d].*?)(?:\s+\d+)?$', text)
        if match:
            chapter_num = match.group(1)
            # 移除末尾的点，保持一致性
            if chapter_num.endswith('.'):
                chapter_num = chapter_num[:-1]
                
            title = match.group(2).strip()
            
            # 任何标题都作为需求项
            chapter_entries.append({
                "chapter": chapter_num,
                "name": f"{chapter_num} {title}",
                "title": title,
                "level": len(chapter_num.split('.'))  # 计算章节级别
            })
    
    print(f"[DEBUG] 段落处理耗时: {time.time() - process_start_time:.2f}秒")
    
    if not chapter_entries:
        print("[WARNING] 目录文件中未找到任何需求条目")
        return []
        
    # 计时过滤子章节
    filter_start_time = time.time()
    
    # 构建章节层级树，用于检测子章节
    chapter_tree = {}
    for entry in chapter_entries:
        chapter_num = entry["chapter"]
        parts = chapter_num.split('.')
        # 构建父章节号 (例如 3.2.1 的父章节为 3.2)
        if len(parts) > 1:
            parent_chapter = '.'.join(parts[:-1])
            if parent_chapter not in chapter_tree:
                chapter_tree[parent_chapter] = []
            chapter_tree[parent_chapter].append(chapter_num)
    
    # 过滤掉有子章节的条目
    filtered_entries = []
    for entry in chapter_entries:
        chapter_num = entry["chapter"]
        if chapter_num not in chapter_tree:  # 如果没有子章节
            filtered_entries.append({
                "name": entry["name"],
                "chapter": entry["chapter"],
                "level": entry["level"]
            })
    
    print(f"[DEBUG] 子章节过滤耗时: {time.time() - filter_start_time:.2f}秒")
    print(f"[DEBUG] 从目录文件中提取了 {len(chapter_entries)} 个条目，过滤后剩余 {len(filtered_entries)} 个需求点")
    
    # 打印前5个和后5个作为示例（如果超过10个的话）
    max_print = min(len(filtered_entries), 5)
    for i, entry in enumerate(filtered_entries[:max_print]):
        print(f"[DEBUG] 需求 {i+1}: {entry['name']} (章节号: {entry['chapter']})")
    
    if len(filtered_entries) > 10:
        print("...")
        
        for i, entry in enumerate(filtered_entries[-max_print:]):
            idx = len(filtered_entries) - max_print + i + 1
            print(f"[DEBUG] 需求 {idx}: {entry['name']} (章节号: {entry['chapter']})")
    
    total_time = time.time() - total_start_time
    print(f"[DEBUG] 需求候选项提取总耗时: {total_time:.2f}秒")
    
    return filtered_entries
