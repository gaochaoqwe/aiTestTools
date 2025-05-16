"""
核心提取模块
实现AI需求提取的主要功能
"""
import os
import time
import json
import logging
from docx import Document

from .api import query_ai_model
from .client import get_client
from .prompts import (
    EXTRACTION_PROMPT_TEMPLATE, 
    CONTINUATION_PROMPT_TEMPLATE,
    SPECIFIC_REQUIREMENT_PROMPT_TEMPLATE,
    SPECIFIC_CONTINUATION_PROMPT_TEMPLATE
)
from .chunking import chunk_document, chunk_document_with_overlap
from .utils import extract_json_from_text, merge_requirement_contents

def ai_extract_requirements(file_path, model=None):
    """
    使用AI提取需求内容，支持滑动窗口和上下文回溯
    
    Args:
        file_path: Word文档路径
        model: 使用的AI模型名称
        
    Returns:
        提取的需求字典 {需求名称: 内容}
    """
    # 记录原始传入的模型名称以方便调试
    logging.debug(f"AI提取需求原始传入模型: {model}")
    
    # 获取AI客户端
    client = get_client()
    if not client:
        logging.error("无法创建AI客户端，请检查配置")
        return {}
    
    # 记录开始时间
    start_time = time.time()
    
    # 从配置文件获取滑动窗口参数
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '../..', 'config.json')
    window_size = 3500  # 默认值
    overlap = 500  # 默认值
    enable_context = True  # 默认启用上下文回溯
    
    if os.path.exists(config_path):
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                extract_params = config.get("extraction_params", {})
                window_size = extract_params.get("window_size", 3500)
                overlap = extract_params.get("overlap", 500)
                enable_context = extract_params.get("enable_context", True)
                logging.info(f"使用配置的滑动窗口参数: 窗口大小={window_size}, 重叠={overlap}, 启用上下文={enable_context}")
        except Exception as e:
            logging.warning(f"读取滑动窗口配置失败: {str(e)}，使用默认配置")
    
    # 加载文档
    doc_load_start = time.time()
    doc = Document(file_path)
    logging.debug(f"文档加载耗时: {time.time() - doc_load_start:.2f}秒")
    
    # 提取文本内容
    text_start_time = time.time()
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    logging.debug(f"文档包含 {len(paragraphs)} 个段落，文本提取耗时: {time.time() - text_start_time:.2f}秒")
    
    # 分块处理文档，使用滑动窗口
    chunk_start_time = time.time()
    if enable_context:
        chunks, last_chapters = chunk_document_with_overlap(paragraphs, window_size, overlap)
        logging.debug(f"文档分成 {len(chunks)} 个滑动窗口，分块耗时: {time.time() - chunk_start_time:.2f}秒")
    else:
        chunks = chunk_document(paragraphs, max_tokens=window_size)  
        last_chapters = [None] * len(chunks)
        logging.debug(f"文档分成 {len(chunks)} 个固定块，分块耗时: {time.time() - chunk_start_time:.2f}秒")
    
    # 处理每个块并收集结果
    all_chapters = []
    last_extracted_chapter = None
    
    for i, (chunk, last_chapter) in enumerate(zip(chunks, last_chapters)):
        logging.info(f"处理块 {i+1}/{len(chunks)}, 长度: {len(chunk)}")
        
        # 基于上下文构建提示
        if i == 0 or not enable_context or not last_extracted_chapter:
            # 第一个块或禁用上下文，使用原始提示
            prompt = EXTRACTION_PROMPT_TEMPLATE.format(chunk)
            logging.debug("使用初始提取提示")
        else:
            # 非第一个块，使用续取提示
            prompt = CONTINUATION_PROMPT_TEMPLATE.format(last_extracted_chapter, chunk)
            logging.info(f"使用续取提示，上一个提取的章节: {last_extracted_chapter}")
        
        try:
            # 查询AI模型
            response = query_ai_model(client, prompt, model)
            
            # 解析JSON响应
            result = extract_json_from_text(response)
            if result and "chapters" in result:
                # 确保 chapters 是列表类型
                chapters = result["chapters"]
                if not isinstance(chapters, list):
                    logging.warning(f"块 {i+1}: chapters 不是列表类型，而是 {type(chapters)}")
                    if isinstance(chapters, dict):
                        # 尝试将单个字典转换为列表
                        chapters = [chapters]
                    else:
                        chapters = []
                
                logging.info(f"块 {i+1}: 成功提取 {len(chapters)} 个章节")
                
                # 更新最后提取的章节，用于下一个块的续取
                if chapters and enable_context:
                    try:
                        last_extracted_chapter = chapters[-1]["chapter_number"] + " " + chapters[-1]["title"]
                        logging.debug(f"更新最后提取的章节为: {last_extracted_chapter}")
                    except (KeyError, IndexError) as e:
                        logging.warning(f"无法更新最后提取的章节: {e}")
                
                all_chapters.extend(chapters)
            else:
                logging.warning(f"块 {i+1}: 未能解析有效的章节数据")
        except Exception as e:
            logging.error(f"处理块 {i+1} 失败: {str(e)}")
    
    # 去除重复章节
    unique_chapters = {}
    for chapter in all_chapters:
        try:
            chapter_key = chapter["chapter_number"] + " " + chapter["title"]
            if chapter_key not in unique_chapters or len(chapter["content"]) > len(unique_chapters[chapter_key]["content"]):
                unique_chapters[chapter_key] = chapter
        except KeyError as e:
            logging.warning(f"章节缺少必要字段: {e}, 跳过此章节: {chapter}")
    
    # 转换为需求字典格式
    requirements = {}
    for chapter_key, chapter in unique_chapters.items():
        if chapter["chapter_number"].startswith("3.1.") or chapter["chapter_number"].startswith("3.2."):
            # 使用章节号+标题作为需求名称
            requirement_name = chapter["title"]
            requirements[requirement_name] = chapter["content"]
    
    logging.debug(f"AI提取完成，耗时: {time.time() - start_time:.2f}秒，提取到 {len(requirements)} 个需求")
    
    return requirements

def ai_extract_specific_requirement(file_path, requirement_name, model=None):
    """
    使用AI提取特定需求的内容，支持滑动窗口处理
    
    Args:
        file_path: Word文档路径
        requirement_name: 需求名称
        model: 使用的AI模型名称
        
    Returns:
        需求内容字符串或None
    """
    # 记录原始传入的模型名称以方便调试
    logging.debug(f"特定AI提取需求原始传入模型: {model}")
    
    # 获取AI客户端
    client = get_client()
    if not client:
        logging.error("无法创建AI客户端，请检查配置")
        return None
    
    # 记录开始时间
    start_time = time.time()
    
    # 从配置文件获取滑动窗口参数
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '../..', 'config.json')
    window_size = 3500  # 默认值
    overlap = 500  # 默认值
    enable_context = True  # 默认启用上下文回溯
    
    if os.path.exists(config_path):
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                extract_params = config.get("extraction_params", {})
                window_size = extract_params.get("window_size", 3500)
                overlap = extract_params.get("overlap", 500)
                enable_context = extract_params.get("enable_context", True)
                logging.info(f"特定需求提取使用配置的滑动窗口参数: 窗口大小={window_size}, 重叠={overlap}, 启用上下文={enable_context}")
        except Exception as e:
            logging.warning(f"读取滑动窗口配置失败: {str(e)}，使用默认配置")
    
    # 加载文档
    doc = Document(file_path)
    logging.debug(f"文档加载耗时: {time.time() - start_time:.2f}秒")
    
    # 提取文本内容
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    logging.debug(f"文档包含 {len(paragraphs)} 个段落")
    
    # 分块处理文档
    if enable_context:
        chunks, _ = chunk_document_with_overlap(paragraphs, window_size, overlap)
        logging.debug(f"文档分成 {len(chunks)} 个滑动窗口")
    else:
        chunks = chunk_document(paragraphs, max_tokens=window_size)
        logging.debug(f"文档分成 {len(chunks)} 个固定块")
    
    # 处理每个块并收集结果
    matched_contents = []
    best_match_score = 0
    
    for i, chunk in enumerate(chunks):
        logging.info(f"处理块 {i+1}/{len(chunks)}, 搜索需求: {requirement_name}")
        
        # 选择使用的提示模板
        if i == 0 or not matched_contents or not enable_context:
            # 第一个块或未找到匹配，使用默认提示
            prompt = SPECIFIC_REQUIREMENT_PROMPT_TEMPLATE.format(requirement_name, chunk)
        else:
            # 已经找到一些内容，使用续取提示
            prompt = SPECIFIC_CONTINUATION_PROMPT_TEMPLATE.format(requirement_name, chunk)
            logging.info("使用续取提示继续搜索需求")
        
        try:
            # 查询AI模型
            response = query_ai_model(client, prompt, model)
            
            # 解析JSON响应
            result = extract_json_from_text(response)
            if result and "requirement" in result and result["requirement"].get("content"):
                content = result["requirement"]["content"]
                
                # 处理content可能是字典的情况
                if isinstance(content, dict):
                    # 如果content是字典，将其转换为字符串
                    logging.info(f"收到的content是字典类型，转换为字符串: {json.dumps(content, ensure_ascii=False)[:100]}...")
                    content = json.dumps(content, ensure_ascii=False, indent=2)
                
                # 确保content是字符串
                if not isinstance(content, str):
                    content = str(content)
                
                # 如果找到的内容提到了需求名称，则存储
                match_score = calculate_relevance_score(content, requirement_name)
                if match_score > 0:
                    matched_contents.append(content)
                    if match_score > best_match_score:
                        best_match_score = match_score
                        logging.info(f"块 {i+1}: 找到更强相关性的匹配，分数: {match_score}")
        except Exception as e:
            logging.error(f"处理块 {i+1} 失败: {str(e)}")
    
    # 合并所有匹配的内容
    final_content = None
    if matched_contents:
        final_content = merge_requirement_contents(matched_contents)
        logging.info(f"合并了 {len(matched_contents)} 个匹配片段，最终内容长度: {len(final_content)}")
    
    logging.debug(f"搜索需求「{requirement_name}」完成，耗时: {time.time() - start_time:.2f}秒")
    
    return final_content

def calculate_relevance_score(content, requirement_name):
    """
    计算内容与需求名称的相关性分数
    
    Args:
        content: 内容文本
        requirement_name: 需求名称
        
    Returns:
        相关性分数 (0-100)
    """
    if not content or not requirement_name:
        return 0
        
    # 转为小写进行比较
    content_lower = content.lower()
    req_lower = requirement_name.lower()
    
    # 基础分：内容包含需求名称
    base_score = 1 if req_lower in content_lower else 0
    if base_score == 0:
        # 检查需求名称的关键词是否存在
        keywords = [word for word in req_lower.split() if len(word) > 3]
        keyword_matches = sum(1 for word in keywords if word in content_lower)
        if keyword_matches > len(keywords) / 2:
            base_score = 0.5
    
    # 字数因子：内容越长，可能包含的信息越多
    length_score = min(1.0, len(content) / 500) if content else 0
    
    # 章节特征：检查是否包含章节号和标题
    structure_score = 0
    if "3.1." in content or "3.2." in content:
        structure_score += 0.5
    if any(heading in content for heading in ["标题", "内容", "依据", "标识号", "说明"]):
        structure_score += 0.5
    
    # 综合评分 (0-100)
    total_score = (base_score * 50) + (length_score * 25) + (structure_score * 25)
    
    return total_score

def ai_extract_named_requirements(file_path, requirement_names, model=None):
    """
    使用AI提取指定需求名称的内容
    
    Args:
        file_path: Word文档路径
        requirement_names: 需求名称列表
        model: 使用的AI模型名称
        
    Returns:
        提取的需求字典 {需求名称: 内容}
    """
    # 记录原始传入的模型名称以方便调试
    logging.debug(f"命名需求提取，原始传入模型: {model}")
    
    result = {}
    start_time = time.time()
    
    for i, req_name in enumerate(requirement_names):
        logging.info(f"提取需求 {i+1}/{len(requirement_names)}: {req_name}")
        content = ai_extract_specific_requirement(file_path, req_name, model)
        if content:
            result[req_name] = content
    
    logging.debug(f"命名需求提取完成，耗时: {time.time() - start_time:.2f}秒，成功提取 {len(result)}/{len(requirement_names)} 个需求")
    
    return result

def ai_rematch_requirements(file_path, matched_requirements, model=None):
    """
    重新匹配未匹配到的需求
    
    Args:
        file_path: Word文档路径
        matched_requirements: 已匹配到的需求字典 {需求名称: 内容}
        model: 使用的AI模型名称
        
    Returns:
        合并后的需求字典 {需求名称: 内容}，包含原有的匹配需求和新匹配的需求
    """
    # 记录原始传入的模型名称以方便调试
    logging.debug(f"AI重新匹配需求原始传入模型: {model}")
    
    # 导入重新匹配模块
    from .rematcher import rematch_requirements
    
    # 执行重新匹配
    new_requirements = rematch_requirements(file_path, matched_requirements, model=model)
    
    # 合并新旧需求
    all_requirements = {**matched_requirements, **new_requirements}
    
    logging.info(f"需求重新匹配完成，原有{len(matched_requirements)}个需求，新增{len(new_requirements)}个需求，总计{len(all_requirements)}个")
    
    return all_requirements
