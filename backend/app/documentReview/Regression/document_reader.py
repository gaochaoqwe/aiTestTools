"""
需求文档读取模块
负责读取Word文档中的需求规格说明
"""
import os
import re
from docx import Document
import glob
import time
import logging

# 设置日志配置
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

def check_word_document(folder_path):
    """
    判断文件夹下是否存在需求规格说明docx文件,如果是就返回对应的路径,如果不是就返回0
    
    参数:
        folder_path: 文件夹路径
        
    返回:
        匹配的文件路径或0
    """
    try:
        # 1、docx文档识别
        pattern = os.path.join(folder_path, "*需求规格说明*.docx")
        matching_files = glob.glob(pattern)
        if matching_files:
            return matching_files[0]
            
        # 2、doc文档识别
        pattern1 = os.path.join(folder_path, "*需求规格说明*.doc")
        matching_files1 = glob.glob(pattern1)
        if matching_files1:
            return matching_files1[0]
            
        # 如果没有匹配到任何文档
        return 0
    except Exception as e:
        print(f"发生异常: {e}")
        return 0





def normalize_chapter_number(text):
    """
    标准化章节号格式，去除多余的空格和尾部的点号
    例如: "3.1.1." -> "3.1.1"，"3. 1. 1" -> "3.1.1"
    
    Args:
        text: 包含章节号的文本
        
    Returns:
        标准化后的章节号
    """
    if not text:
        return ""
    
    # 先用正则提取章节号部分
    match = re.search(r'(\d+\.\d+\.\d+\.?)', text)
    if not match:
        return text
        
    chapter = match.group(1)
    # 移除所有空格
    chapter = chapter.replace(" ", "")
    # 如果以点号结尾，移除点号
    if chapter.endswith('.'):
        chapter = chapter[:-1]
    
    return chapter

def is_table_of_contents(paragraph, nearby_paragraphs=None):
    """
    判断段落是否属于目录页面
    
    Args:
        paragraph: 当前段落文本
        nearby_paragraphs: 附近段落列表，用于上下文判断
        
    Returns:
        是否为目录页面的布尔值
    """
    # 如果段落包含明确的需求内容标记，不应被视为目录
    if re.search(r'(新增需求描述|修改设计描述|依据：|a\s*\)|标识号：|说明：)', paragraph):
        return False
    
    # 目录特征1: 包含页码指示（数字+空格）在行末
    has_page_numbers = bool(re.search(r'\s\d+\s*$', paragraph))
    
    # 目录特征2: 包含大量的点(.)或空格连接符
    dots_count = paragraph.count('.')
    dots_ratio = dots_count / len(paragraph) if paragraph else 0
    has_many_dots = dots_ratio > 0.2  # 提高阈值，避免误判
    
    # 目录特征3: 短行，且包含章节编号模式，且以页码结尾
    is_short_line = len(paragraph) < 60  # 减小长度限制，只匹配非常短的行
    has_section_number = bool(re.search(r'^\s*\d+\.\d+\.', paragraph))
    ends_with_page_number = bool(re.search(r'\s\d+\s*$', paragraph))
    
    # 如果有附近段落用于上下文判断
    if nearby_paragraphs:
        # 检查附近段落是否都符合目录特征
        nearby_features = []
        for p in nearby_paragraphs:
            # 递归调用但不传递上下文，避免无限递归
            if p != paragraph:  # 避免自引用
                # 简化判断，只检查基本目录特征
                page_num = bool(re.search(r'\s\d+\s*$', p))
                dots = p.count('.') / len(p) if p else 0 > 0.2
                is_short = len(p) < 60
                nearby_features.append(page_num and (dots or is_short))
        
        nearby_toc_ratio = sum(nearby_features) / len(nearby_features) if nearby_features else 0
        
        # 提高上下文要求：附近段落必须有更高比例是目录
        if nearby_toc_ratio > 0.8:  # 提高阈值
            return is_short_line and (has_page_numbers or has_many_dots)
    
    # 需满足更严格的组合条件才视为目录
    return is_short_line and has_page_numbers and (has_many_dots or has_section_number)

def identify_toc_sections(paragraphs):
    """
    识别文档中的目录部分
    
    Args:
        paragraphs: 文档段落列表
        
    Returns:
        目录段落的索引集合
    """
    toc_indices = set()
    toc_section_start = -1
    in_toc_section = False
    consecutive_toc_count = 0
    
    # 第一步：找出明确的目录页
    # 目录通常出现在文档前半部分，所以只检查前30%的段落
    search_limit = min(len(paragraphs), int(len(paragraphs) * 0.3))
    
    for i in range(search_limit):
        para = paragraphs[i]
        nearby = paragraphs[max(0, i-2):i] + paragraphs[i+1:min(len(paragraphs), i+3)]
        
        if is_table_of_contents(para, nearby):
            consecutive_toc_count += 1
            if consecutive_toc_count > 4:  # 需要连续5个以上段落都是目录特征
                if not in_toc_section:
                    toc_section_start = i - consecutive_toc_count + 1
                    in_toc_section = True
                toc_indices.add(i)
        else:
            # 如果连续3个段落都不是目录，视为目录段结束
            if consecutive_toc_count > 0:
                consecutive_toc_count -= 1  # 逐渐减少计数，允许目录中有少量非目录段落
            else:
                if in_toc_section:
                    # 结束当前目录区域
                    in_toc_section = False
    
    # 修复：确认起止位置合理后再添加这个区间的所有段落
    if toc_section_start >= 0 and len(toc_indices) > 10:  # 至少要有10个确认的目录条目
        # 扩展目录区域
        toc_section_end = max(toc_indices) + 1
        # 将这个范围内的段落都标记为目录
        toc_indices.update(range(toc_section_start, toc_section_end))
    
    logging.debug(f"识别到可能的目录段落: {len(toc_indices)} 个")
    return toc_indices

def is_real_requirement_section(paragraphs, start_index, normalized_chapter=None):
    """
    判断指定段落是否是真正的需求内容（而非目录或其他内容）
    通过检查结构特征（依据、说明、对关联编码的影响等）来判断
    
    Args:
        paragraphs: 文档的所有段落
        start_index: 当前段落在文档中的索引
        normalized_chapter: 标准化的章节号，用于匹配
        
    Returns:
        是否为真正需求内容的布尔值
    """
    # 如果索引超出范围，返回False
    if start_index >= len(paragraphs):
        return False
        
    # 获取当前段落及后续段落
    current_para = paragraphs[start_index]
    following_paras = paragraphs[start_index+1:min(start_index+20, len(paragraphs))]
    
    # 拼接所有段落为一个文本，用于整体特征匹配
    combined_text = "\n".join([current_para] + following_paras)
    
    # 特征1：包含"依据"和"说明"关键词
    has_basis = bool(re.search(r'依据[：:]\s*', combined_text))
    has_description = bool(re.search(r'说明[：:]\s*', combined_text))
    
    # 特征2：包含"标识号"和字母数字组合的标识
    has_identifier = bool(re.search(r'标识号[：:]\s*\w+', combined_text))
    
    # 特征3：包含"对关联xxx的影响"结构
    has_impact_analysis = bool(re.search(r'对关联(需求|设计|编码)的影响', combined_text))
    
    # 特征4：包含子章节结构（如3.1.1.1）
    has_sub_sections = False
    if normalized_chapter:
        sub_chapter_pattern = f"{normalized_chapter}\\.\\d+"
        has_sub_sections = bool(re.search(sub_chapter_pattern, combined_text))
    
    # 特征5：包含a)、b)、c)等列表项目
    has_list_items = bool(re.search(r'[a-e]\s*\)\s*', combined_text))
    
    # 特征组合判断：至少满足3个特征，或者同时有依据和影响分析
    score = sum([has_basis, has_description, has_identifier, has_impact_analysis, has_sub_sections, has_list_items])
    return score >= 3 or (has_basis and has_impact_analysis)

def read_xuqiu_wendang_document(file_path, requirement_names=None, catalog_file_path=None):
    """
    读取需求文档，提取需求内容
    
    Args:
        file_path: 文档路径
        requirement_names: 待提取的需求名称，如果为None则提取所有需求
        catalog_file_path: 目录文件路径，可选
        
    Returns:
        需求内容字典 {需求名称: 内容}
    """
    # 记录开始时间，用于性能统计
    start_time = time.time()
    
    # 加载文档
    doc = Document(file_path)
    logging.debug(f"文档加载耗时: {time.time() - start_time:.2f}秒")
    
    # 提取文本内容
    text_start_time = time.time()
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    logging.debug(f"文档包含 {len(paragraphs)} 个段落，文本提取耗时: {time.time() - text_start_time:.2f}秒")
    
    # 识别目录部分，避免匹配到目录中的需求名称
    toc_detection_start = time.time()
    toc_indices = identify_toc_sections(paragraphs)
    logging.debug(f"目录检测耗时: {time.time() - toc_detection_start:.2f}秒，识别到 {len(toc_indices)} 个目录段落")
    
    # 打印前10个段落样例，用于调试
    for i in range(min(10, len(paragraphs))):
        logging.debug(f"段落 {i} 示例: {paragraphs[i]}"+ (" [目录]" if i in toc_indices else ""))
    
    if not requirement_names:
        # 如果没有指定需求名称，则尝试从文档中提取
        requirement_names = extract_requirement_names(paragraphs)
    
    logging.debug(f"read_xuqiu_wendang_document - 需要提取的需求项数量: {len(requirement_names)}")
    
    # 查找每个需求在文档中的位置
    requirement_positions = []
    for req_name in requirement_names:
        logging.debug(f"寻找需求「{req_name}」的位置...")
        found = False
        
        # 使用纯文本匹配，不依赖章节号
        # 先在非目录部分搜索需求名称的关键词
        key_terms = req_name.split()
        # 取长度超过1个字符的词，避免匹配太短的通用词
        key_terms = [term for term in key_terms if len(term) > 1]
        
        # 如果需求名称很短，使用整个需求名称
        if not key_terms:
            key_terms = [req_name]
            
        logging.debug(f"从需求「{req_name}」提取的关键词: {key_terms}")
        
        # 匹配可能的候选段落
        candidates = []
        
        # 遍历所有段落寻找包含关键词的部分
        for i, para in enumerate(paragraphs):
            # 跳过目录部分
            if i in toc_indices:
                continue
                
            # 计算匹配分数：包含多少个关键词
            match_score = sum(1 for term in key_terms if term in para)
            
            # 如果包含任何关键词，添加到候选列表
            if match_score > 0:
                # 检查是否包含需求特征
                has_features = "依据" in para or "标识号" in para or "说明" in para
                
                candidates.append({
                    "index": i,
                    "score": match_score,
                    "has_features": has_features,
                    "text": para[:50] + "..."  # 保存前50个字符用于调试
                })
                
        # 按匹配分数和需求特征排序
        candidates.sort(key=lambda x: (x["score"], 1 if x["has_features"] else 0), reverse=True)
        
        # 输出前三个候选段落用于调试
        for i, candidate in enumerate(candidates[:3]):
            logging.debug(f"候选段落 {candidate['index']} 匹配分数: {candidate['score']}, 包含需求特征: {candidate['has_features']}, 内容: {candidate['text']}...")
        
        # 选择最佳匹配
        if candidates:
            best_match = candidates[0]
            logging.debug(f"选择最佳匹配: 段落 {best_match['index']}, 分数 {best_match['score']}, 包含需求特征: {best_match['has_features']}, 内容: {best_match['text']}...")
            requirement_positions.append((req_name, best_match["index"]))
            found = True
        
        if not found:
            logging.warning(f"未找到需求「{req_name}」的位置")
    
    # 如果没有找到任何需求位置，返回空字典
    if not requirement_positions:
        logging.warning("未找到任何需求位置")
        return {}
    
    # 对需求位置按照在文档中的顺序排序
    requirement_positions.sort(key=lambda x: x[1])
    logging.debug(f"排序后的需求位置: {[(name, pos) for name, pos in requirement_positions]}")
    
    # 收集所有需求内容
    requirement_contents = {}
    
    # 预扫描文档，提取所有可能的需求名称作为分隔标志
    requirement_names_for_separation = extract_requirement_names(paragraphs)
    logging.debug(f"检测到 {len(requirement_names_for_separation)} 个可能的需求名称标志")
    
    # 遍历所有需求位置，提取内容
    for i, (req_name, position) in enumerate(requirement_positions):
        # 确定内容结束位置
        if i < len(requirement_positions) - 1:
            # 如果不是最后一个需求，结束位置是下一个需求的开始位置
            end_position = requirement_positions[i + 1][1]
        else:
            # 如果是最后一个需求，结束位置是文档末尾
            # 但为了避免包含无关内容，限制为开始位置后的30行或文档末尾
            end_position = min(position + 30, len(paragraphs))
        
        # 提取内容
        content_paragraphs = extract_requirement_content(paragraphs, position, end_position, toc_indices)
        
        # 如果内容段落数太少，可能是匹配到了错误的位置
        if len(content_paragraphs) < 2:
            logging.warning(f"需求「{req_name}」的内容段落太少（{len(content_paragraphs)}），可能是错误的匹配")
            continue
        
        # 合并段落为一个字符串
        content = "\n".join(content_paragraphs)
        
        logging.debug(f"需求「{req_name}」内容提取成功，长度: {len(content)} 字符 ，包含 {len(content_paragraphs)} 个段落")
        
        # 添加到结果字典
        requirement_contents[req_name] = content
    
    logging.debug(f"提取到{len(requirement_contents)}个需求内容")
    for name, content in requirement_contents.items():
        logging.debug(f"需求「{name}」的内容长度: {len(content)} 字符")
    
    return requirement_contents

def extract_requirement_content(paragraphs, start_pos, end_pos, toc_indices):
    """
    根据起始位置和结束位置，提取需求内容段落
    
    Args:
        paragraphs: 文档段落列表
        start_pos: 需求开始位置
        end_pos: 需求结束位置
        toc_indices: 目录段落索引集合
        
    Returns:
        需求内容段落列表
    """
    content_paragraphs = []
    
    # 首先确认起始段落是否包含需求特征
    has_req_features = False
    for i in range(start_pos, min(start_pos + 5, end_pos)):
        if "依据" in paragraphs[i] or "标识号" in paragraphs[i] or "说明" in paragraphs[i]:
            has_req_features = True
            break
    
    # 如果起始位置附近有需求特征，则从这开始提取
    # 否则从start_pos开始
    current_pos = start_pos
    
    # 提前获取所有可能的需求名称，用于检测下一个需求的开始
    next_requirement_indicators = []
    for para in paragraphs:
        # 典型需求名称标志，比如"3.2.3 重构激光雷达数据处理算法"
        if re.match(r"^\d+\.\d+\.\d+\s+[\u4e00-\u9fa5]+", para):
            # 去掉章节号，只保留需求名称部分
            name_part = re.sub(r"^\d+\.\d+\.\d+\s+", "", para)
            if len(name_part) > 3:  # 排除太短的匹配
                next_requirement_indicators.append(name_part)
    
    logging.debug(f"检测到 {len(next_requirement_indicators)} 个可能的需求名称标志")
    
    # 收集内容直到结束位置或者找到分隔符
    while current_pos < end_pos:
        # 跳过目录段落
        if current_pos in toc_indices:
            current_pos += 1
            continue
            
        # 当前段落
        current_para = paragraphs[current_pos]
        
        # 检查该段落是否标志着下一个需求的开始
        is_next_requirement = False
        
        # 检查1: 是否包含标准需求起始特征
        if current_pos > start_pos + 3 and (("依据" in current_para or "标识号" in current_para) and len(current_para) < 100):
            # 这可能是下一个需求的开始
            is_next_requirement = True
            logging.debug(f"在段落 {current_pos} 检测到下一个需求的特征标志: {current_para[:30]}...")
            
        # 检查2: 是否匹配其他需求的名称
        if not is_next_requirement and next_requirement_indicators:
            for indicator in next_requirement_indicators:
                if indicator in current_para and len(indicator) > 5 and current_pos > start_pos + 3:
                    is_next_requirement = True
                    logging.debug(f"在段落 {current_pos} 检测到匹配其他需求名称的内容: {indicator}")
                    break
        
        # 检查3: 章节格式检查
        if not is_next_requirement and current_pos > start_pos + 3:
            # 检查是否是一个新的主章节标题 (比如 "3.2.3 重构激光雷达...") 
            # 章节标题通常较短
            if re.match(r"^\d+\.\d+\.\d+\s+[\u4e00-\u9fa5]", current_para) and len(current_para) < 50:
                is_next_requirement = True
                logging.debug(f"在段落 {current_pos} 检测到新的章节标题: {current_para[:30]}...")
        
        # 如果检测到下一个需求，并且我们已经收集了足够的内容，就停止
        if is_next_requirement and len(content_paragraphs) > 3:
            logging.debug(f"在段落 {current_pos} 检测到下一个需求开始，停止提取")
            break
        
        # 添加当前段落
        content_paragraphs.append(current_para)
        
        # 移动到下一段
        current_pos += 1
    
    # 检查提取的内容是否过长，如果是则进一步分析结构截断
    if len(content_paragraphs) > 10:
        for i in range(5, len(content_paragraphs)):
            para = content_paragraphs[i]
            # 如果发现疑似下一个需求开始的结构标记，截断到这里
            if re.match(r"^\d+\.\d+\.\d+\.\d+\.", para) or "修改文件" in para or "修改前" in para:
                logging.debug(f"基于内容结构在第 {i} 段截断，内容: {para[:30]}...")
                content_paragraphs = content_paragraphs[:i]
                break
    
    return content_paragraphs

def extract_requirement_names(paragraphs):
    """
    从文档段落中提取可能的需求名称
    
    Args:
        paragraphs: 文档段落列表
    
    Returns:
        需求名称列表
    """
    requirement_names = []
    
    # 匹配像 "3.1.1." 或 "3.1.1" 这样的章节号格式，后面通常跟着需求名称
    pattern = re.compile(r'^(\d+\.\d+\.\d+\.?\s*)(.*?)$')
    
    for para in paragraphs:
        match = pattern.match(para)
        if match:
            section_number = match.group(1).strip()
            section_title = match.group(2).strip()
            
            # 清理section_number确保格式一致（移除末尾的点）
            section_number = normalize_chapter_number(section_number)
            
            # 提取的需求名称格式为 "3.1.1 需求标题"
            requirement_name = f"{section_number} {section_title}"
            requirement_names.append(requirement_name)
            logging.debug(f"从段落提取到需求: {requirement_name}")
    
    logging.debug(f"从文档中提取到 {len(requirement_names)} 个可能的需求名称")
    return requirement_names

def extract_requirement_candidates(document_file_path, catalog_file_path):
    """
    从目录文件提取需求名称和章节号。
    
    参数:
        document_file_path: str, 主文档文件路径
        catalog_file_path: str, 目录文件路径
        
    目录文件格式应该类似图片中所示，包含章节号和需求名称（页码不是必需的）。
    如果章节下有子章节，则过滤掉该章节。
    返回：包含需求名称和章节号的对象列表 [{name, chapter}, ...]
    """
    from docx import Document
    import re
    import time

    # 设置环境变量，以便其他函数能够访问当前文档路径
    import os
    os.environ['CURRENT_DOC_PATH'] = document_file_path
    
    if not catalog_file_path:
        print("[ERROR] 必须提供目录文件！")
        return []
        
    print(f"[DEBUG] 使用目录文件提取需求点: {catalog_file_path}")
    
    total_start_time = time.time()
    
    # 从目录文件读取内容
    try:
        doc_load_start = time.time()
        doc = Document(catalog_file_path)
        print(f"[DEBUG] 目录文件加载耗时: {time.time() - doc_load_start:.2f}秒")
    except Exception as e:
        print(f"[ERROR] 打开目录文件时出错: {str(e)}")
        return []
    
    # 存储章节号和名称
    chapter_entries = []
    
    # 计时提取段落
    para_start_time = time.time()
    paragraphs = [para for para in doc.paragraphs if para.text.strip()]
    print(f"[DEBUG] 目录文件包含 {len(paragraphs)} 个非空段落，提取耗时: {time.time() - para_start_time:.2f}秒")
    
    # 计时处理段落
    process_start_time = time.time()
    
    for i, para in enumerate(paragraphs):
        if i % 50 == 0:  # 每处理50个段落输出一次进度
            print(f"[DEBUG] 正在处理段落 {i+1}/{len(paragraphs)}...")
            
        text = para.text.strip()
        if not text:
            continue
            
        # 匹配章节号和名称
        # 格式如: "3.2.1. 用户身份验证功能    2" 或 "3.2.1 用户身份验证功能    2" 或 "3.2.1 用户身份验证功能"
        # 章节号可以有多个层级 (如 3.2.1.4.5)，页码是可选的
        match = re.match(r'^(\d+(?:\.\d+)*(?:\.)?)\s+([^\d].*?)(?:\s+\d+)?$', text)
        if match:
            chapter_num = match.group(1)
            # 移除末尾的点，保持一致性
            if chapter_num.endswith('.'):
                chapter_num = chapter_num[:-1]
                
            title = match.group(2).strip()
            
            # 任何标题都作为需求项
            chapter_entries.append({
                "chapter": chapter_num,
                "name": title,  # 修改：只使用标题作为需求名称，不包含章节号
                "title": title,
                "level": len(chapter_num.split('.'))  # 计算章节级别
            })
    
    print(f"[DEBUG] 段落处理耗时: {time.time() - process_start_time:.2f}秒")
    
    if not chapter_entries:
        print("[WARNING] 目录文件中未找到任何需求条目")
        return []
        
    # 计时过滤子章节
    filter_start_time = time.time()
    
    # 构建章节层级树，用于检测子章节
    chapter_tree = {}
    for entry in chapter_entries:
        chapter_num = entry["chapter"]
        parts = chapter_num.split('.')
        # 构建父章节号 (例如 3.2.1 的父章节为 3.2)
        if len(parts) > 1:
            parent_chapter = '.'.join(parts[:-1])
            if parent_chapter not in chapter_tree:
                chapter_tree[parent_chapter] = []
            chapter_tree[parent_chapter].append(chapter_num)
    
    # 过滤掉有子章节的条目
    filtered_entries = []
    for entry in chapter_entries:
        chapter_num = entry["chapter"]
        if chapter_num not in chapter_tree:  # 如果没有子章节
            filtered_entries.append({
                "name": entry["title"],  # 修改：只使用标题作为需求名称，不包含章节号
                "chapter": entry["chapter"],
                "level": entry["level"]
            })
    
    print(f"[DEBUG] 子章节过滤耗时: {time.time() - filter_start_time:.2f}秒")
    print(f"[DEBUG] 从目录文件中提取了 {len(chapter_entries)} 个条目，过滤后剩余 {len(filtered_entries)} 个需求点")
    
    # 打印前5个和后5个作为示例（如果超过10个的话）
    max_print = min(len(filtered_entries), 5)
    for i, entry in enumerate(filtered_entries[:max_print]):
        print(f"[DEBUG] 需求 {i+1}: {entry['name']} (章节号: {entry['chapter']})")
    
    if len(filtered_entries) > 10:
        print("...")
        
        for i, entry in enumerate(filtered_entries[-max_print:]):
            idx = len(filtered_entries) - max_print + i + 1
            print(f"[DEBUG] 需求 {idx}: {entry['name']} (章节号: {entry['chapter']})")
    
    total_time = time.time() - total_start_time
    print(f"[DEBUG] 需求候选项提取总耗时: {total_time:.2f}秒")
    
    return filtered_entries
